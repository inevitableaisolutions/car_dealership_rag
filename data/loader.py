"""Document loading utilities for multiple file formats."""

import json
import csv
from pathlib import Path
from typing import List
from io import BytesIO, StringIO

from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as DocxDocument


def load_pdf(file_path: str = None, file_bytes: bytes = None) -> List[Document]:
    """Load text from a PDF file."""
    if file_bytes:
        reader = PdfReader(BytesIO(file_bytes))
    else:
        reader = PdfReader(file_path)
    
    documents = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text.strip():
            documents.append(Document(
                page_content=text,
                metadata={"source": file_path or "uploaded_pdf", "page": i + 1}
            ))
    
    return documents


def load_docx(file_path: str = None, file_bytes: bytes = None) -> List[Document]:
    """Load text from a Word document."""
    if file_bytes:
        doc = DocxDocument(BytesIO(file_bytes))
    else:
        doc = DocxDocument(file_path)
    
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))
    
    if full_text:
        return [Document(
            page_content="\n\n".join(full_text),
            metadata={"source": file_path or "uploaded_docx"}
        )]
    return []


def load_json(file_path: str = None, file_bytes: bytes = None) -> List[Document]:
    """Load documents from a JSON file."""
    if file_bytes:
        data = json.loads(file_bytes.decode("utf-8"))
    else:
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    
    documents = []
    
    # Handle different JSON structures
    if isinstance(data, list):
        for i, item in enumerate(data):
            if isinstance(item, dict):
                # Convert dict to readable text
                content = format_dict_as_text(item)
                documents.append(Document(
                    page_content=content,
                    metadata={"source": file_path or "uploaded_json", "index": i}
                ))
            else:
                documents.append(Document(
                    page_content=str(item),
                    metadata={"source": file_path or "uploaded_json", "index": i}
                ))
    elif isinstance(data, dict):
        content = format_dict_as_text(data)
        documents.append(Document(
            page_content=content,
            metadata={"source": file_path or "uploaded_json"}
        ))
    
    return documents


def load_csv(file_path: str = None, file_bytes: bytes = None) -> List[Document]:
    """Load documents from a CSV file."""
    if file_bytes:
        content = file_bytes.decode("utf-8")
        reader = csv.DictReader(StringIO(content))
    else:
        with open(file_path, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
    
    documents = []
    for i, row in enumerate(reader):
        content = format_dict_as_text(dict(row))
        documents.append(Document(
            page_content=content,
            metadata={"source": file_path or "uploaded_csv", "row": i + 1}
        ))
    
    return documents


def load_txt(file_path: str = None, file_bytes: bytes = None) -> List[Document]:
    """Load text from a plain text file."""
    if file_bytes:
        content = file_bytes.decode("utf-8")
    else:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
    
    if content.strip():
        return [Document(
            page_content=content,
            metadata={"source": file_path or "uploaded_txt"}
        )]
    return []


def format_dict_as_text(data: dict) -> str:
    """Format a dictionary as readable text."""
    lines = []
    for key, value in data.items():
        # Clean up key name
        clean_key = key.replace("_", " ").title()
        lines.append(f"{clean_key}: {value}")
    return "\n".join(lines)


def load_file(file_path: str = None, file_bytes: bytes = None, file_name: str = None) -> List[Document]:
    """
    Load documents from a file based on its extension.
    
    Args:
        file_path: Path to the file (for local files)
        file_bytes: Raw file bytes (for uploaded files)
        file_name: Original filename (needed when using file_bytes)
    """
    if file_path:
        extension = Path(file_path).suffix.lower()
    elif file_name:
        extension = Path(file_name).suffix.lower()
    else:
        raise ValueError("Either file_path or file_name must be provided")
    
    loaders = {
        ".pdf": load_pdf,
        ".docx": load_docx,
        ".doc": load_docx,
        ".json": load_json,
        ".csv": load_csv,
        ".txt": load_txt,
    }
    
    loader = loaders.get(extension)
    if not loader:
        raise ValueError(f"Unsupported file format: {extension}")
    
    return loader(file_path=file_path, file_bytes=file_bytes)
